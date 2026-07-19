import sys
sys.path.insert(0, r'D:\Do_aN\.codex_tmp\pydeps')
from docx import Document

doc = Document(r'D:\Do_aN\Report3_Software Requirement Specification.docx')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith(('4.1 ', '4.2 ', '4.3 ', '5. ')) or t in ('4.2 Mobile Screen', '4.3 Web Screen'):
        print(i, repr(t[:120]), p.style.name)
print('paragraphs', len(doc.paragraphs), 'tables', len(doc.tables), 'sections', len(doc.sections))
