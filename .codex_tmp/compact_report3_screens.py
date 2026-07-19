import re
import sys
from pathlib import Path

sys.path.insert(0, r'D:\Do_aN\.codex_tmp\pydeps')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = Path(r'D:\Do_aN\Report3_Software Requirement Specification.docx')


def el_text(el):
    return ''.join(el.itertext()).strip()


def clear_paragraph(p):
    for child in list(p._p):
        if child.tag != qn('w:pPr'):
            p._p.remove(child)


def set_labeled_paragraph(p, label, value):
    clear_paragraph(p)
    r1 = p.add_run(f'{label}:')
    r1.bold = True
    p.add_run(f' {value}')


def new_labeled_before(doc, anchor_el, label, value):
    p = doc.add_paragraph(style='normal')
    set_labeled_paragraph(p, label, value)
    body = doc._element.body
    body.remove(p._p)
    anchor_el.addprevious(p._p)
    return p


def applies_to(platform):
    mapping = {
        'Common Mobile App': (
            'Mother Mobile App, Family Member Mobile App, and Verified Expert Mobile App '
            'where the flow is available.'
        ),
        'Shared Mobile Apps': (
            'Mother Mobile App, Family Member Mobile App, and Verified Expert Mobile App '
            'where the flow is available.'
        ),
        'Mother Mobile App': 'Mother Mobile App.',
        'Family Member Mobile App': 'Family Member Mobile App.',
        'Verified Expert App': 'Verified Expert Mobile App.',
        'Common Web Portal': (
            'Verified Expert Web Portal, Moderator Web Portal, Content Admin Web Portal, '
            'and System Admin Web Portal where the flow is available.'
        ),
        'Shared Web Portals': (
            'Verified Expert Web Portal, Moderator Web Portal, Content Admin Web Portal, '
            'and System Admin Web Portal where the flow is available.'
        ),
        'Expert Web Portal': 'Verified Expert Web Portal.',
        'Moderator Web Portal': 'Moderator Web Portal.',
        'Content Admin Web Portal': 'Content Admin Web Portal.',
        'System Admin Web Portal': 'System Admin Web Portal.',
    }
    if platform not in mapping:
        raise ValueError(f'Unknown platform: {platform!r}')
    return mapping[platform]


def display_platform(platform):
    return {
        'Common Mobile App': 'Shared Mobile Apps',
        'Common Web Portal': 'Shared Web Portals',
        'Verified Expert App': 'Verified Expert Mobile App',
    }.get(platform, platform)


doc = Document(DOCX)
body = doc._element.body
paragraphs = doc.paragraphs
start_p = next(p for p in paragraphs if p.text.strip() == '4.2 Mobile Screen')
end_p = next(p for p in paragraphs if p.text.strip() == '5. Non-Functional Requirements')
children = list(body)
start_idx = children.index(start_p._p)
end_idx = children.index(end_p._p)
section = children[start_idx:end_idx]

screen_headings = [
    el for el in section
    if el.tag == qn('w:p') and re.match(r'^4\.[23]\.\d+\.\d+\s+', el_text(el))
]

removed_prompt_tables = 0
updated_field_tables = 0

for index, heading_el in enumerate(screen_headings):
    all_children = list(body)
    h_idx = all_children.index(heading_el)
    next_heading_el = screen_headings[index + 1] if index + 1 < len(screen_headings) else end_p._p
    n_idx = all_children.index(next_heading_el)
    block = all_children[h_idx + 1:n_idx]

    meta_el = next(el for el in block if el.tag == qn('w:p') and el_text(el).startswith('Platform:'))
    purpose_el = next(el for el in block if el.tag == qn('w:p') and el_text(el).startswith('Purpose:'))
    meta_text = el_text(meta_el)
    m = re.match(r'^Platform:\s*(.*?)\s+Feature:\s*(.*?)\s+Release:\s*', meta_text)
    if not m:
        raise ValueError(f'Cannot parse metadata: {meta_text!r}')
    platform, feature = m.group(1).strip(), m.group(2).strip()
    purpose_text = el_text(purpose_el).split(':', 1)[1].strip()

    meta_p = next(p for p in doc.paragraphs if p._p is meta_el)
    purpose_p = next(p for p in doc.paragraphs if p._p is purpose_el)
    new_labeled_before(doc, meta_el, 'Applies To', applies_to(platform))
    set_labeled_paragraph(meta_p, 'Platform', display_platform(platform))
    new_labeled_before(doc, purpose_el, 'Feature', feature)
    set_labeled_paragraph(purpose_p, 'Purpose', purpose_text)

    # Re-read the block after inserting the compact metadata paragraphs.
    all_children = list(body)
    h_idx = all_children.index(heading_el)
    n_idx = all_children.index(next_heading_el)
    block = all_children[h_idx + 1:n_idx]

    prompt_label = next((el for el in block if el.tag == qn('w:p') and el_text(el) == 'Stitch Wireframe Prompt'), None)
    figure_caption = next((el for el in block if el.tag == qn('w:p') and el_text(el).startswith('Figure 4.')), None)
    tables = [el for el in block if el.tag == qn('w:tbl')]
    if len(tables) != 2:
        raise ValueError(f'Expected two tables for {el_text(heading_el)!r}, found {len(tables)}')
    prompt_table, field_table = tables

    if prompt_label is not None:
        body.remove(prompt_label)
    body.remove(prompt_table)
    removed_prompt_tables += 1
    if figure_caption is not None:
        body.remove(figure_caption)

    # Only the header row is bold. All body-cell text is regular.
    table_obj = next(t for t in doc.tables if t._tbl is field_table)
    for r_idx, row in enumerate(table_obj.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = (r_idx == 0)
    updated_field_tables += 1

settings = doc.settings._element
update_fields = settings.find(qn('w:updateFields'))
if update_fields is None:
    update_fields = OxmlElement('w:updateFields')
    settings.append(update_fields)
update_fields.set(qn('w:val'), 'true')

doc.save(DOCX)
print(f'screens={len(screen_headings)} prompt_tables_removed={removed_prompt_tables} field_tables_updated={updated_field_tables}')
