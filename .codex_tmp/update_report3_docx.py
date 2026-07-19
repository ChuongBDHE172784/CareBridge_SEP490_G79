import re
import sys
from pathlib import Path

sys.path.insert(0, r'D:\Do_aN\.codex_tmp\pydeps')

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Twips

ROOT = Path(r'D:\Do_aN')
DOCX = ROOT / 'Report3_Software Requirement Specification.docx'
MD = ROOT / 'CareBridge_SRS_4.2_Mobile_4.3_Web_Screens_Release1_121UC_Updated.docx.md'


def clean_md(text: str) -> str:
    text = text.strip()
    text = re.sub(r'^\*\*(.*?)\*\*$', r'\1', text)
    text = text.replace(r'\#', '#').replace(r'\&', '&').replace(r'\-', '-')
    text = text.replace(r'\_', '_').replace(r'\*', '*')
    return text


def split_table_row(line: str):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [clean_md(cell) for cell in re.split(r'(?<!\\)\|', line)]


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement('w:tblHeader')
    marker.set(qn('w:val'), 'true')
    tr_pr.append(marker)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width))
    tc_w.set(qn('w:type'), 'dxa')


def format_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = borders.find(qn(f'w:{edge}'))
        if border is None:
            border = OxmlElement(f'w:{edge}')
            borders.append(border)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '777777')
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = tc_pr.find(qn('w:tcMar'))
            if mar is None:
                mar = OxmlElement('w:tcMar')
                tc_pr.append(mar)
            for edge in ('top', 'left', 'bottom', 'right'):
                node = mar.find(qn(f'w:{edge}'))
                if node is None:
                    node = OxmlElement(f'w:{edge}')
                    mar.append(node)
                node.set(qn('w:w'), '100')
                node.set(qn('w:type'), 'dxa')


def add_markdown_paragraph(doc, line):
    bold_whole = bool(re.fullmatch(r'\*\*.+\*\*', line.strip()))
    p = doc.add_paragraph(style='normal')
    run = p.add_run(clean_md(line))
    run.bold = bold_whole
    if line.startswith(('Figure ', '**Figure ')) or line.startswith(('Table ', '**Table ')):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.bold = True
    return p


def add_markdown_section(doc, text, usable_width):
    lines = text.splitlines()
    added = []
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue
        heading = re.match(r'^(#{1,3})\s+\*\*(.*?)\*\*\s*$', stripped)
        if heading:
            level = len(heading.group(1)) + 2
            p = doc.add_paragraph(clean_md(heading.group(2)), style=f'Heading {level}')
            added.append(p._p)
            i += 1
            continue
        if stripped.startswith('|'):
            block = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                block.append(lines[i].strip())
                i += 1
            rows = [split_table_row(x) for x in block]
            rows = [row for row in rows if not all(re.fullmatch(r':?-{2,}:?', c.replace(' ', '')) for c in row)]
            if not rows:
                continue
            cols = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            if cols == 1:
                widths = [usable_width]
            elif cols == 2:
                widths = [int(usable_width * 0.32), usable_width - int(usable_width * 0.32)]
            else:
                widths = [usable_width // cols] * cols
                widths[-1] += usable_width - sum(widths)
            for r_idx, values in enumerate(rows):
                values += [''] * (cols - len(values))
                for c_idx, value in enumerate(values):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = value
                    for p in cell.paragraphs:
                        p.style = doc.styles['normal']
                        if r_idx == 0 and cols > 1 or c_idx == 0 and cols > 1:
                            for run in p.runs:
                                run.bold = True
            if cols > 1:
                set_repeat_table_header(table.rows[0])
            format_table(table, widths)
            added.append(table._tbl)
            continue
        p = add_markdown_paragraph(doc, stripped)
        added.append(p._p)
        i += 1
    return added


doc = Document(DOCX)
body = doc._element.body

start_p = next(p for p in doc.paragraphs if p.text.strip() == '4.2 Mobile Screen')
end_p = next(p for p in doc.paragraphs if p.text.strip() == '5. Non-Functional Requirements')
start_el = start_p._p
end_el = end_p._p

children = list(body)
start_idx = children.index(start_el)
end_idx = children.index(end_el)
for el in children[start_idx:end_idx]:
    body.remove(el)

source = MD.read_text(encoding='utf-8')
source_start = source.index('# **4.2 Mobile Screen**')
release_text = source[source_start:].strip()

section = doc.sections[0]
usable_width = int((section.page_width - section.left_margin - section.right_margin) / 635)
new_elements = add_markdown_section(doc, release_text, usable_width)

for el in new_elements:
    body.remove(el)
    end_el.addprevious(el)

settings = doc.settings._element
update_fields = settings.find(qn('w:updateFields'))
if update_fields is None:
    update_fields = OxmlElement('w:updateFields')
    settings.append(update_fields)
update_fields.set(qn('w:val'), 'true')

doc.save(DOCX)
print(f'updated={DOCX}')
print(f'new_elements={len(new_elements)} usable_width={usable_width}')
