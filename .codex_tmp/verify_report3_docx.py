import re
import sys
sys.path.insert(0, r'D:\Do_aN\.codex_tmp\pydeps')
from docx import Document

path = r'D:\Do_aN\Report3_Software Requirement Specification.docx'
doc = Document(path)
paras = doc.paragraphs
start = next(i for i,p in enumerate(paras) if p.text.strip() == '4.2 Mobile Screen')
web = next(i for i,p in enumerate(paras) if p.text.strip() == '4.3 Web Screen')
end = next(i for i,p in enumerate(paras) if p.text.strip() == '5. Non-Functional Requirements')
screen_re = re.compile(r'^4\.[23]\.\d+\.\d+\s+')
mobile = [p.text.strip() for p in paras[start:web] if screen_re.match(p.text.strip())]
webs = [p.text.strip() for p in paras[web:end] if screen_re.match(p.text.strip())]
body = doc._element.body
els = list(body)
s_el = paras[start]._p
e_el = paras[end]._p
section_els = els[els.index(s_el):els.index(e_el)]
tables = [el for el in section_els if el.tag.endswith('}tbl')]
print('paragraphs', len(paras), 'all_tables', len(doc.tables), 'sections', len(doc.sections))
print('mobile', len(mobile), 'web', len(webs), 'total', len(mobile)+len(webs))
print('section_tables', len(tables), 'start', start, 'web_start', web, 'end', end)
print('first', mobile[0], 'last', webs[-1])
print('forbidden', [p.text for p in paras[start:end] if any(x in p.text for x in ('Payment Screen','Realtime Consultation Session Screen','Partner Portal Landing Screen'))])
assert len(mobile) == 61
assert len(webs) == 18
assert len(tables) == 158
assert paras[end].style.name == 'Heading 2'
