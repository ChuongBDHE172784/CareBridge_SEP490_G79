import sys
sys.path.insert(0, r'D:\Do_aN\.codex_tmp\pydeps')
from docx import Document

path = r'D:\Do_aN\Report3_Software Requirement Specification.docx'
doc = Document(path)


def clear_paragraph(p):
    from docx.oxml.ns import qn
    for child in list(p._p):
        if child.tag != qn('w:pPr'):
            p._p.remove(child)


def set_labeled(p, label, value):
    clear_paragraph(p)
    r = p.add_run(label + ':')
    r.bold = True
    p.add_run(' ' + value)


def collapse_exact_repetition(value):
    for count in range(5, 1, -1):
        if len(value) % count == 0:
            unit = value[:len(value) // count]
            if unit * count == value:
                return unit
    return value


removed = 0
fixed = 0
for p in list(doc.paragraphs):
    if p.text.strip() == 'Stitch Wireframe Prompt':
        p._p.getparent().remove(p._p)
        removed += 1
    elif p.text.startswith('Purpose:'):
        value = p.text.split(':', 1)[1].strip()
        clean = value.split('Purpose:', 1)[0].strip()
        clean = collapse_exact_repetition(clean)
        if clean != value:
            set_labeled(p, 'Purpose', clean)
            fixed += 1

doc.save(path)
print('stitch_labels_removed', removed, 'purposes_fixed', fixed)
