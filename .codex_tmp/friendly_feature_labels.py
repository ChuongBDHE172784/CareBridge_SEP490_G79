import sys
sys.path.insert(0, r'D:\Do_aN\.codex_tmp\pydeps')
from docx import Document
from docx.oxml.ns import qn

path = r'D:\Do_aN\Report3_Software Requirement Specification.docx'
doc = Document(path)

mapping = {
    'MF-01 Account, Trust & Access Control': 'Authentication & setup',
    'MF-02 Mother Care Journey': 'Mother care journey',
    'MF-03 Baby Care Journey, Growth and Vaccination': 'Baby care journey, growth & vaccination',
    'MF-04 Community Q&A and Moderation': 'Community Q&A & moderation',
    'MF-04 Community Q&A and Moderation / MF-05 Expert Contribution': 'Community Q&A & expert contribution',
    'MF-05 Verified Expert Network and Contribution': 'Verified expert network & contribution',
    'MF-05 Verified Expert Network and Contribution / MF-07 Emergency Map': 'Verified expert network & nearby support',
    'MF-06 AI Nurse Assistant and Risk Triage': 'AI nurse assistant & risk triage',
    'MF-07 Emergency Map and Nearby Care Support': 'Emergency map & nearby care support',
    'MF-08 Personal Health Records and Source Labeling': 'Personal health records & source labeling',
    'MF-09 Reminders, Tasks and Care Plan': 'Reminders, tasks & care plan',
    'MF-10 Family Sync and Cooperative Care': 'Family sync & cooperative care',
    'MF-11 Verified Content and Checklist Hub': 'Verified content & checklist hub',
    'MF-12 Expense and Preparation Planner': 'Expense & preparation planner',
    'MF-13 Connected Device and Health Data Integration': 'Connected device & health data integration',
    'MF-14 Smart Activity Monitoring and Safety Support': 'Smart activity monitoring & safety support',
}

count = 0
for p in doc.paragraphs:
    if not p.text.startswith('Feature:'):
        continue
    old = p.text.split(':', 1)[1].strip()
    new = mapping[old]
    for child in list(p._p):
        if child.tag != qn('w:pPr'):
            p._p.remove(child)
    r = p.add_run('Feature:')
    r.bold = True
    p.add_run(' ' + new)
    count += 1

doc.save(path)
print('feature_labels_updated', count)
