import re
import sys
sys.path.insert(0, r'D:\Do_aN\.codex_tmp\pydeps')
from docx import Document
from docx.oxml.ns import qn

path = r'D:\Do_aN\Report3_Software Requirement Specification.docx'
doc = Document(path)
body = doc._element.body
paras = doc.paragraphs
start = next(p for p in paras if p.text.strip() == '4.2 Mobile Screen')
end = next(p for p in paras if p.text.strip() == '5. Non-Functional Requirements')
els = list(body)
section = els[els.index(start._p):els.index(end._p)]
text = '\n'.join(p.text for p in paras[paras.index(start):paras.index(end)])
headings = [el for el in section if el.tag == qn('w:p') and re.match(r'^4\.[23]\.\d+\.\d+\s+', ''.join(el.itertext()).strip())]
tables = [el for el in section if el.tag == qn('w:tbl')]
labels = {name: 0 for name in ('Applies To', 'Platform', 'Feature', 'Purpose')}
label_errors = []
for p in paras[paras.index(start):paras.index(end)]:
    for label in labels:
        if p.text.startswith(label + ':'):
            labels[label] += 1
            bold_text = ''.join(r.text for r in p.runs if r.bold)
            if bold_text != label + ':':
                label_errors.append((p.text, bold_text))

table_errors = []
for t in doc.tables:
    if t._tbl not in tables:
        continue
    if len(t.columns) != 2 or t.cell(0,0).text.strip() != 'Field name' or t.cell(0,1).text.strip() != 'Description':
        table_errors.append('bad header')
        continue
    for r_idx, row in enumerate(t.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if bool(run.bold) != (r_idx == 0):
                        table_errors.append(f'row {r_idx}: {run.text!r} bold={run.bold}')

print('screens', len(headings), 'tables', len(tables), 'labels', labels)
print('stitch_count', text.count('Stitch Wireframe Prompt'), 'figure_count', len(re.findall(r'^Figure 4\.', text, re.M)))
print('label_errors', len(label_errors), 'table_errors', len(table_errors))
sample_start = next(i for i,p in enumerate(paras) if p.text.strip() == '4.2.1.1 Mobile Welcome Screen')
for p in paras[sample_start:sample_start+5]:
    print(repr(p.text), [(r.text, r.bold) for r in p.runs])
assert len(headings) == 79
assert len(tables) == 79
assert all(v == 79 for v in labels.values())
assert not label_errors
assert not table_errors
assert 'Stitch Wireframe Prompt' not in text
assert not re.search(r'^Figure 4\.', text, re.M)
